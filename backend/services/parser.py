
import os
from openai import AzureOpenAI
from backend.core.config import settings
import json
from pypdf import PdfReader
from docx import Document
import docx2txt
from io import BytesIO

class DocumentParserService:
    def __init__(self):
        self.client = None
        if settings.AZURE_OPENAI_API_KEY and settings.AZURE_OPENAI_ENDPOINT:
            self.client = AzureOpenAI(
                api_key=settings.AZURE_OPENAI_API_KEY,
                api_version=settings.AZURE_OPENAI_API_VERSION,
                azure_endpoint=settings.AZURE_OPENAI_ENDPOINT
            )

    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """
        Extract text from PDF bytes using pypdf.
        """
        try:
            reader = PdfReader(BytesIO(file_content))
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            print(f"PDF Extraction Error: {e}")
            return ""

    def extract_text_from_docx(self, file_content: bytes) -> str:
        """
        Extract text from DOCX bytes using python-docx.
        """
        try:
            # Method 1: Using python-docx (better for structured content)
            doc = Document(BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            return text
        except Exception as e:
            print(f"DOCX Extraction Error (python-docx): {e}")
            # Fallback to docx2txt
            try:
                text = docx2txt.process(BytesIO(file_content))
                return text
            except Exception as e2:
                print(f"DOCX Extraction Error (docx2txt): {e2}")
                return ""

    def extract_text_from_doc(self, file_content: bytes) -> str:
        """
        Extract text from DOC (legacy Word format).
        Note: .doc format is complex and requires external tools.
        This is a basic implementation that may not work for all .doc files.
        """
        try:
            # Try to use docx2txt (works for some .doc files)
            text = docx2txt.process(BytesIO(file_content))
            if text:
                return text

            # For production, consider using:
            # - LibreOffice (via subprocess)
            # - antiword (via subprocess)
            # - Online conversion service
            print("Warning: .doc format support is limited. Consider converting to .docx for best results.")
            return ""
        except Exception as e:
            print(f"DOC Extraction Error: {e}")
            print("Note: .doc format requires external tools. Please convert to .docx or .pdf")
            return ""

    def extract_text(self, file_content: bytes, file_type: str) -> str:
        """
        Universal text extraction method that handles multiple formats.
        """
        file_type = file_type.lower()

        if file_type == "pdf":
            return self.extract_text_from_pdf(file_content)
        elif file_type == "docx":
            return self.extract_text_from_docx(file_content)
        elif file_type == "doc":
            return self.extract_text_from_doc(file_content)
        else:
            print(f"Unsupported file type: {file_type}")
            return ""

    def analyze_text_with_llm(self, text: str) -> list[dict]:
        """
        Analyze text using Azure OpenAI to extract deadlines and events.
        """
        if not self.client:
            print("Azure OpenAI client not initialized.")
            return []

        # Truncate text if too long (simple approach for MVP)
        # Ideally we should chunk it, but for most docs < 50 pages, full context is better if token limit allows.
        # GPT-4o has 128k context, so usually fine.
        truncated_text = text[:100000] 

        prompt = """
        You are an expert legal document analyzer. Your task is to extract all deadline events, payment dates, deliverables, and important milestones from the provided text.

        IMPORTANT: Keep the ORIGINAL LANGUAGE of the document. If the document is in Chinese, use Chinese for titles and descriptions. If in English, use English. DO NOT translate.

        Return the result as a JSON list of objects. Each object should have:
        - "title": A short descriptive title of the event in the SAME LANGUAGE as the source document (string)
        - "due_date": The deadline date in ISO 8601 format (YYYY-MM-DD) or null if relative/unknown (string or null)
        - "description": Context or details about the event in the SAME LANGUAGE as the source document (string)
        - "confidence_score": Your confidence in this extraction (0-100 integer)
        - "source_text": The snippet from the text where you found this information (string)

        If no events are found, return an empty list.
        Just return the JSON array, no markdown formatting or other text.
        """

        try:
            response = self.client.chat.completions.create(
                model=settings.AZURE_OPENAI_DEPLOYMENT,
                messages=[
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": f"Here is the document text:\n\n{truncated_text}"}
                ],
                temperature=0,
                response_format={"type": "json_object"} # or just standard if model doesn't support json_object mode perfectly yet, but gpt-4 usually does. Azure deployment gpt-4.1 might act like gpt-4-turbo.
            )
            
            content = response.choices[0].message.content
            # Parse JSON
            # The model might return {"events": [...]} or just [...]
            # We asked for a list, but sometimes it wraps.
            data = json.loads(content)
            
            if isinstance(data, list):
                return data
            elif isinstance(data, dict) and "events" in data:
                return data["events"]
            elif isinstance(data, dict):
                 # Try to find any list in the dict values
                 for val in data.values():
                     if isinstance(val, list):
                         return val
            return []

        except Exception as e:
            print(f"LLM Analysis Error: {e}")
            return []

parser_service = DocumentParserService()
